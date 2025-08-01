import os
import json
import pandas as pd
import numpy as np
from scipy import stats
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.utils import PlotlyJSONEncoder
import io
import base64
from flask import Flask, render_template, request, jsonify, send_file, flash, redirect, url_for
from flask_cors import CORS
from werkzeug.utils import secure_filename
import openai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import PyPDF2
import tempfile
import re
from datetime import datetime
from control_charts import ControlChartAnalyzer
from advanced_statistics import AdvancedStatisticalAnalyzer

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)
app.secret_key = os.environ.get('SECRET_KEY', 'autostatiq-secret-key-2024')

# Configuration
UPLOAD_FOLDER = 'uploads'
RESULTS_FOLDER = 'results'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'csv', 'xlsx', 'xls', 'docx', 'json'}

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(RESULTS_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# OpenAI API configuration
openai.api_key = os.environ.get('OPENAI_API_KEY')

class StatisticalAnalyzer:
    def __init__(self):
        self.data = None
        self.results = {}
        self.control_chart_analyzer = ControlChartAnalyzer()
        self.advanced_analyzer = AdvancedStatisticalAnalyzer()
        
    def load_data(self, file_path, file_type):
        """Load data from various file formats"""
        try:
            if file_type in ['csv']:
                self.data = pd.read_csv(file_path)
            elif file_type in ['xlsx', 'xls']:
                self.data = pd.read_excel(file_path)
            elif file_type == 'json':
                self.data = pd.read_json(file_path)
            elif file_type == 'txt':
                # Try to read as CSV first, then as plain text
                try:
                    self.data = pd.read_csv(file_path, sep='\t')
                except:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    return content
            elif file_type == 'docx':
                return self.extract_text_from_docx(file_path)
            elif file_type == 'pdf':
                return self.extract_text_from_pdf(file_path)
            
            return self.data
        except Exception as e:
            raise Exception(f"Error loading data: {str(e)}")
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    
    def is_structured_data(self, data):
        """Determine if the data is structured (DataFrame) or unstructured text"""
        return isinstance(data, pd.DataFrame)
    
    def analyze_with_gpt(self, question_or_text):
        """Analyze natural language questions or text with GPT-4"""
        try:
            prompt = f"""
            You are a statistical analysis expert. Analyze the following text/question and provide:
            1. What type of statistical analysis is needed
            2. What specific statistical tests should be performed
            3. If this contains data, extract and structure it
            4. Provide clear recommendations for analysis
            
            Text/Question: {question_or_text}
            
            Please respond in JSON format with the following structure:
            {{
                "analysis_type": "descriptive/inferential/predictive",
                "recommended_tests": ["test1", "test2"],
                "data_structure": "description of data if present",
                "recommendations": "specific analysis recommendations"
            }}
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            return {
                "analysis_type": "descriptive", 
                "recommended_tests": ["descriptive_stats"],
                "data_structure": "text analysis",
                "recommendations": f"Error in GPT analysis: {str(e)}"
            }
    
    def perform_descriptive_analysis(self, data):
        """Perform descriptive statistical analysis"""
        results = {}
        
        # Basic statistics
        numeric_cols = data.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) > 0:
            results['descriptive_stats'] = data[numeric_cols].describe().to_dict()
            
            # Correlation matrix
            if len(numeric_cols) > 1:
                results['correlation_matrix'] = data[numeric_cols].corr().to_dict()
            
            # Generate visualizations
            results['plots'] = self.generate_plots(data, numeric_cols)
        
        # Categorical analysis
        categorical_cols = data.select_dtypes(include=['object']).columns
        if len(categorical_cols) > 0:
            results['categorical_analysis'] = {}
            for col in categorical_cols:
                results['categorical_analysis'][col] = data[col].value_counts().to_dict()
        
        return results
    
    def perform_inferential_tests(self, data, test_type='auto'):
        """Perform inferential statistical tests"""
        results = {}
        numeric_cols = data.select_dtypes(include=[np.number]).columns
        
        if len(numeric_cols) >= 2:
            # T-test for two numeric columns
            col1, col2 = numeric_cols[0], numeric_cols[1]
            t_stat, p_value = stats.ttest_ind(data[col1].dropna(), data[col2].dropna())
            results['t_test'] = {
                'statistic': float(t_stat),
                'p_value': float(p_value),
                'columns': [col1, col2]
            }
            
            # Correlation test
            corr_coef, corr_p = stats.pearsonr(data[col1].dropna(), data[col2].dropna())
            results['correlation_test'] = {
                'correlation_coefficient': float(corr_coef),
                'p_value': float(corr_p),
                'columns': [col1, col2]
            }
        
        # ANOVA if we have categorical and numeric data
        categorical_cols = data.select_dtypes(include=['object']).columns
        if len(categorical_cols) > 0 and len(numeric_cols) > 0:
            cat_col = categorical_cols[0]
            num_col = numeric_cols[0]
            groups = [group[num_col].dropna() for name, group in data.groupby(cat_col)]
            if len(groups) > 1:
                f_stat, p_value = stats.f_oneway(*groups)
                results['anova'] = {
                    'f_statistic': float(f_stat),
                    'p_value': float(p_value),
                    'categorical_column': cat_col,
                    'numeric_column': num_col
                }
        
        return results
    
    def perform_regression_analysis(self, data):
        """Perform regression analysis"""
        results = {}
        numeric_cols = data.select_dtypes(include=[np.number]).columns
        
        if len(numeric_cols) >= 2:
            # Simple linear regression with first two numeric columns
            X = data[numeric_cols[0]].dropna().values.reshape(-1, 1)
            y = data[numeric_cols[1]].dropna().values
            
            # Ensure X and y have the same length
            min_len = min(len(X), len(y))
            X, y = X[:min_len], y[:min_len]
            
            model = LinearRegression()
            model.fit(X, y)
            
            predictions = model.predict(X)
            r2 = r2_score(y, predictions)
            
            results['linear_regression'] = {
                'r_squared': float(r2),
                'coefficient': float(model.coef_[0]),
                'intercept': float(model.intercept_),
                'independent_variable': numeric_cols[0],
                'dependent_variable': numeric_cols[1]
            }
        
        return results
    
    def generate_plots(self, data, numeric_cols):
        """Generate visualization plots"""
        plots = {}
        
        # Histogram for first numeric column
        if len(numeric_cols) > 0:
            fig = px.histogram(data, x=numeric_cols[0], title=f'Distribution of {numeric_cols[0]}')
            plots['histogram'] = json.dumps(fig, cls=PlotlyJSONEncoder)
        
        # Scatter plot for two numeric columns
        if len(numeric_cols) >= 2:
            fig = px.scatter(data, x=numeric_cols[0], y=numeric_cols[1], 
                           title=f'{numeric_cols[0]} vs {numeric_cols[1]}')
            plots['scatter'] = json.dumps(fig, cls=PlotlyJSONEncoder)
        
        # Correlation heatmap
        if len(numeric_cols) > 2:
            corr_matrix = data[numeric_cols].corr()
            fig = px.imshow(corr_matrix, text_auto=True, aspect="auto", 
                          title="Correlation Heatmap")
            plots['heatmap'] = json.dumps(fig, cls=PlotlyJSONEncoder)
        
        return plots
    
    def perform_control_chart_analysis(self, data, question=""):
        """Perform Statistical Process Control (SPC) analysis with control charts"""
        try:
            # Check if SPC analysis is needed
            recommendations = self.control_chart_analyzer.detect_control_chart_needs(data, question)
            
            if not recommendations['recommended_charts']:
                return {}
            
            control_chart_results = {
                'spc_analysis': True,
                'recommendations': recommendations,
                'control_charts': {},
                'interpretations': {}
            }
            
            # Generate control charts based on recommendations
            for chart_rec in recommendations['recommended_charts']:
                chart_type = chart_rec['type']
                column = chart_rec['column']
                
                try:
                    stats_serializable = None
                    chart_b64 = None
                    
                    if chart_type == 'i_chart':
                        # Individual chart
                        data_values = data[column].dropna().tolist()
                        if len(data_values) >= 2:
                            fig, stats = self.control_chart_analyzer.create_i_chart(
                                data_values, f"Individual Chart - {column}"
                            )
                            chart_b64 = self.control_chart_analyzer.save_chart_as_base64(fig)
                            stats_serializable = self._convert_numpy_types(stats)
                    
                    elif chart_type == 'mr_chart':
                        # Moving Range chart
                        data_values = data[column].dropna().tolist()
                        if len(data_values) >= 2:
                            fig, stats = self.control_chart_analyzer.create_mr_chart(
                                data_values, f"Moving Range Chart - {column}"
                            )
                            chart_b64 = self.control_chart_analyzer.save_chart_as_base64(fig)
                            stats_serializable = self._convert_numpy_types(stats)
                    
                    elif chart_type == 'c_chart':
                        # C chart for count of defects
                        data_values = data[column].dropna().astype(int).tolist()
                        if len(data_values) >= 5:
                            fig, stats = self.control_chart_analyzer.create_c_chart(
                                data_values, f"C Chart - {column}"
                            )
                            chart_b64 = self.control_chart_analyzer.save_chart_as_base64(fig)
                            stats_serializable = self._convert_numpy_types(stats)
                    
                    # Add chart if successfully created
                    if chart_b64 and stats_serializable:
                        chart_key = f'{chart_type}_{column}'
                        control_chart_results['control_charts'][chart_key] = {
                            'chart_image': chart_b64,
                            'statistics': stats_serializable,
                            'type': chart_rec.get('description', f'{chart_type.upper()} Chart'),
                            'column': column
                        }
                        
                        # Add interpretation
                        interpretation = self.control_chart_analyzer.generate_control_chart_interpretation(
                            chart_type, stats_serializable
                        )
                        control_chart_results['interpretations'][chart_key] = interpretation
                
                except Exception as chart_error:
                    print(f"Error creating {chart_type} for {column}: {str(chart_error)}")
                    continue
            
            return control_chart_results
            
        except Exception as e:
            print(f"Error in control chart analysis: {str(e)}")
            return {}
    
    def perform_advanced_statistical_analysis(self, data, user_question=""):
        """Perform comprehensive advanced statistical analysis using all available methods"""
        try:
            advanced_results = {}
            
            # Comprehensive descriptive analysis
            desc_results = self.advanced_analyzer.perform_comprehensive_descriptive_analysis(data)
            if desc_results:
                advanced_results.update(desc_results)
            
            # Frequency analysis for categorical variables
            freq_results = self.advanced_analyzer.perform_frequency_analysis(data)
            if freq_results:
                advanced_results.update(freq_results)
            
            # Cross-tabulation analysis
            cross_tab_results = self.advanced_analyzer.perform_cross_tabulation_analysis(data)
            if cross_tab_results:
                advanced_results.update(cross_tab_results)
            
            # Advanced correlation analysis
            corr_results = self.advanced_analyzer.perform_advanced_correlation_analysis(data)
            if corr_results:
                advanced_results.update(corr_results)
            
            # Regression analysis
            reg_results = self.advanced_analyzer.perform_advanced_regression_analysis(data)
            if reg_results:
                advanced_results.update(reg_results)
            
            # Logistic regression for binary outcomes
            log_reg_results = self.advanced_analyzer.perform_logistic_regression_analysis(data)
            if log_reg_results:
                advanced_results.update(log_reg_results)
            
            # Comprehensive hypothesis testing
            hyp_results = self.advanced_analyzer.perform_comprehensive_hypothesis_testing(data, user_question)
            if hyp_results:
                advanced_results.update(hyp_results)
            
            # Advanced ANOVA with post-hoc tests
            anova_results = self.advanced_analyzer.perform_advanced_anova(data)
            if anova_results:
                advanced_results.update(anova_results)
            
            # Normality tests
            norm_results = self.advanced_analyzer.perform_normality_tests(data)
            if norm_results:
                advanced_results.update(norm_results)
            
            # Outlier detection
            outlier_results = self.advanced_analyzer.perform_outlier_detection(data)
            if outlier_results:
                advanced_results.update(outlier_results)
            
            # Missing data analysis
            missing_results = self.advanced_analyzer.perform_missing_data_analysis(data)
            if missing_results:
                advanced_results.update(missing_results)
            
            # Principal Component Analysis
            pca_results = self.advanced_analyzer.perform_principal_component_analysis(data)
            if pca_results:
                advanced_results.update(pca_results)
            
            # Clustering analysis
            cluster_results = self.advanced_analyzer.perform_clustering_analysis(data)
            if cluster_results:
                advanced_results.update(cluster_results)
            
            # Convert numpy types for JSON serialization
            advanced_results = self.advanced_analyzer._convert_numpy_types(advanced_results)
            
            return advanced_results
            
        except Exception as e:
            print(f"Error in advanced statistical analysis: {str(e)}")
            return {}
    
    def _convert_numpy_types(self, obj):
        """Convert numpy types to Python native types for JSON serialization"""
        import numpy as np
        
        if isinstance(obj, dict):
            return {key: self._convert_numpy_types(value) for key, value in obj.items()}
        elif isinstance(obj, list):
            return [self._convert_numpy_types(item) for item in obj]
        elif isinstance(obj, np.integer):
            return int(obj)
        elif isinstance(obj, np.floating):
            return float(obj)
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        else:
            return obj
    
    def get_gpt_interpretation(self, results):
        """Get GPT interpretation of statistical results"""
        try:
            prompt = f"""
            As a statistical expert, please interpret the following statistical analysis results:
            
            {json.dumps(results, indent=2)}
            
            Please provide:
            1. A clear interpretation of each statistical test
            2. What the results mean in practical terms
            3. Any limitations or caveats
            4. Recommendations for further analysis
            
            Keep the explanation accessible to non-statisticians while being technically accurate.
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            
            return response.choices[0].message.content or "No interpretation available"
        except Exception as e:
            return f"Error generating interpretation: {str(e)}"
    
    def get_gpt_conclusion(self, results, interpretation):
        """Get GPT conclusion and recommendations"""
        try:
            prompt = f"""
            Based on the statistical analysis results and interpretation below, provide a clear conclusion and actionable recommendations:
            
            Results: {json.dumps(results, indent=2)}
            
            Interpretation: {interpretation}
            
            Please provide:
            1. A concise executive summary
            2. Key findings
            3. Practical recommendations
            4. Next steps for further analysis if needed
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            
            return response.choices[0].message.content or "No conclusion available"
        except Exception as e:
            return f"Error generating conclusion: {str(e)}"

class ReportGenerator:
    def __init__(self):
        self.doc = Document()
        
    def create_report(self, analysis_results, interpretation, conclusion, original_question="", dataset_info=""):
        """Create a professional DOCX report"""
        
        # Title
        title = self.doc.add_heading('AutoStatIQ Statistical Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = self.doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
        
        # Executive Summary
        self.doc.add_heading('Executive Summary', level=1)
        self.doc.add_paragraph(conclusion)
        
        # Original Question/Dataset
        if original_question:
            self.doc.add_heading('Research Question', level=1)
            self.doc.add_paragraph(original_question)
        
        if dataset_info:
            self.doc.add_heading('Dataset Information', level=1)
            self.doc.add_paragraph(dataset_info)
        
        # Statistical Results
        self.doc.add_heading('Statistical Analysis Results', level=1)
        self.add_results_tables(analysis_results)
        
        # Interpretation
        self.doc.add_heading('Interpretation', level=1)
        self.doc.add_paragraph(interpretation)
        
        # Footer
        self.doc.add_page_break()
        footer_para = self.doc.add_paragraph('© Roman Chaudhary | Contact: chaudharyroman.com.np')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return self.doc
    
    def add_results_tables(self, results):
        """Add statistical results as tables"""
        for key, value in results.items():
            if key == 'plots':
                continue  # Skip plots for DOCX
            elif key == 'control_charts':
                self.add_control_charts_section(value, results.get('interpretations', {}))
                continue
            elif key == 'interpretations' and 'control_charts' in results:
                continue  # Skip interpretations if control charts exist (already handled)
                
            self.doc.add_heading(key.replace('_', ' ').title(), level=2)
            
            if isinstance(value, dict):
                # Create table for dictionary data
                if key == 'descriptive_stats':
                    self.add_descriptive_stats_table(value)
                elif key == 'correlation_matrix':
                    self.add_correlation_table(value)
                elif key == 'spc_analysis':
                    continue  # Skip boolean flag
                elif key == 'recommendations':
                    self.add_spc_recommendations_section(value)
                else:
                    # Generic key-value table
                    table = self.doc.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Header
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Metric'
                    hdr_cells[1].text = 'Value'
                    
                    # Data rows
                    for k, v in value.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(k)
                        row_cells[1].text = str(v)
            else:
                self.doc.add_paragraph(str(value))
    
    def add_control_charts_section(self, control_charts, interpretations):
        """Add control charts section to report"""
        if not control_charts:
            return
        
        self.doc.add_heading('Statistical Process Control (SPC) Analysis', level=1)
        self.doc.add_paragraph(
            "Control charts are used to monitor process stability and detect "
            "special causes of variation. The following charts have been generated "
            "based on your data characteristics:"
        )
        
        for chart_key, chart_data in control_charts.items():
            chart_type = chart_data.get('type', 'Control Chart')
            column = chart_data.get('column', 'Unknown')
            
            # Chart heading
            self.doc.add_heading(f"{chart_type} - {column}", level=2)
            
            # Add chart statistics table
            if 'statistics' in chart_data:
                stats = chart_data['statistics']
                table = self.doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Control Chart Parameter'
                hdr_cells[1].text = 'Value'
                
                # Add key statistics
                key_stats = ['center_line', 'ucl', 'lcl']
                for stat in key_stats:
                    if stat in stats:
                        row_cells = table.add_row().cells
                        stat_name = stat.replace('_', ' ').title()
                        if stat == 'ucl':
                            stat_name = 'Upper Control Limit'
                        elif stat == 'lcl':
                            stat_name = 'Lower Control Limit'
                        elif stat == 'center_line':
                            stat_name = 'Center Line'
                        
                        row_cells[0].text = stat_name
                        row_cells[1].text = f"{stats[stat]:.4f}"
                
                # Add out-of-control points info
                ooc_points = stats.get('out_of_control_points', [])
                if ooc_points:
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Out-of-Control Points'
                    row_cells[1].text = f"{len(ooc_points)} points detected"
                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = 'Process Status'
                    row_cells[1].text = 'In Control'
            
            # Add interpretation
            if chart_key in interpretations:
                self.doc.add_paragraph()
                self.doc.add_paragraph(interpretations[chart_key])
            
            self.doc.add_paragraph()
    
    def add_spc_recommendations_section(self, recommendations):
        """Add SPC recommendations section"""
        if not recommendations or not recommendations.get('recommended_charts'):
            return
        
        self.doc.add_paragraph("Based on the data analysis, the following control charts are recommended:")
        
        for i, chart_rec in enumerate(recommendations['recommended_charts'], 1):
            chart_type = chart_rec.get('type', 'Unknown')
            description = chart_rec.get('description', 'Control chart analysis')
            column = chart_rec.get('column', 'data')
            
            self.doc.add_paragraph(f"{i}. {description} for column '{column}'")
        
        if recommendations.get('reasoning'):
            self.doc.add_paragraph("Reasoning:")
            for reason in recommendations['reasoning']:
                self.doc.add_paragraph(f"• {reason}")
    
    def add_descriptive_stats_table(self, desc_stats):
        """Add descriptive statistics table"""
        if not desc_stats:
            return
            
        # Get column names
        columns = list(desc_stats.keys())
        stats_names = list(desc_stats[columns[0]].keys()) if columns else []
        
        table = self.doc.add_table(rows=len(stats_names) + 1, cols=len(columns) + 1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        table.rows[0].cells[0].text = 'Statistic'
        for i, col in enumerate(columns):
            table.rows[0].cells[i + 1].text = col
        
        # Data rows
        for i, stat in enumerate(stats_names):
            table.rows[i + 1].cells[0].text = stat
            for j, col in enumerate(columns):
                value = desc_stats[col][stat]
                table.rows[i + 1].cells[j + 1].text = f"{value:.4f}" if isinstance(value, float) else str(value)
    
    def add_correlation_table(self, corr_matrix):
        """Add correlation matrix table"""
        if not corr_matrix:
            return
            
        columns = list(corr_matrix.keys())
        
        table = self.doc.add_table(rows=len(columns) + 1, cols=len(columns) + 1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row and column
        table.rows[0].cells[0].text = ''
        for i, col in enumerate(columns):
            table.rows[0].cells[i + 1].text = col
            table.rows[i + 1].cells[0].text = col
        
        # Data
        for i, row_col in enumerate(columns):
            for j, col_col in enumerate(columns):
                value = corr_matrix[row_col][col_col]
                table.rows[i + 1].cells[j + 1].text = f"{value:.4f}"

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    """Main dashboard page"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and analysis"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file selected'}), 400
        
        file = request.files['file']
        question = request.form.get('question', '')
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Get file extension
            file_ext = filename.rsplit('.', 1)[1].lower()
            
            # Initialize analyzer
            analyzer = StatisticalAnalyzer()
            
            # Load data
            data = analyzer.load_data(file_path, file_ext)
            
            # Determine analysis type
            if analyzer.is_structured_data(data):
                # Structured data analysis
                results = {}
                
                # Perform different types of analysis
                results.update(analyzer.perform_descriptive_analysis(data))
                results.update(analyzer.perform_inferential_tests(data))
                results.update(analyzer.perform_regression_analysis(data))
                
                # Perform control chart analysis if applicable
                spc_results = analyzer.perform_control_chart_analysis(data, question)
                if spc_results:
                    results.update(spc_results)
                
                # Perform advanced statistical analysis
                advanced_results = analyzer.perform_advanced_statistical_analysis(data, question)
                if advanced_results:
                    results.update(advanced_results)
                
                # Get GPT interpretation and conclusion
                interpretation = analyzer.get_gpt_interpretation(results)
                conclusion = analyzer.get_gpt_conclusion(results, interpretation)
                
                dataset_info = f"Dataset contains {len(data)} rows and {len(data.columns)} columns.\nColumns: {', '.join(data.columns.tolist())}"
                
            else:
                # Unstructured text analysis
                gpt_analysis = analyzer.analyze_with_gpt(data)
                
                results = {
                    'text_analysis': gpt_analysis,
                    'content_length': len(data),
                    'analysis_type': gpt_analysis.get('analysis_type', 'text analysis')
                }
                
                interpretation = analyzer.get_gpt_interpretation(results)
                conclusion = analyzer.get_gpt_conclusion(results, interpretation)
                
                dataset_info = f"Text document with {len(data)} characters analyzed."
            
            # Generate report
            report_gen = ReportGenerator()
            doc = report_gen.create_report(
                results, interpretation, conclusion, 
                question, dataset_info
            )
            
            # Save report
            report_filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_path = os.path.join(RESULTS_FOLDER, report_filename)
            doc.save(report_path)
            
            # Clean up uploaded file
            os.remove(file_path)
            
            return jsonify({
                'success': True,
                'message': 'Analysis completed successfully',
                'results': results,
                'interpretation': interpretation,
                'conclusion': conclusion,
                'report_filename': report_filename,
                'plots': results.get('plots', {})
            })
        
        else:
            return jsonify({'error': 'File type not allowed'}), 400
    
    except Exception as e:
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/download/<filename>')
def download_report(filename):
    """Download generated report"""
    try:
        file_path = os.path.join(RESULTS_FOLDER, filename)
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

@app.route('/api/analyze', methods=['POST'])
def api_analyze():
    """REST API endpoint for analysis"""
    try:
        data = request.get_json()
        
        if 'question' in data:
            # Text analysis
            analyzer = StatisticalAnalyzer()
            gpt_analysis = analyzer.analyze_with_gpt(data['question'])
            
            results = {
                'text_analysis': gpt_analysis,
                'question': data['question']
            }
            
            interpretation = analyzer.get_gpt_interpretation(results)
            conclusion = analyzer.get_gpt_conclusion(results, interpretation)
            
            return jsonify({
                'success': True,
                'results': results,
                'interpretation': interpretation,
                'conclusion': conclusion
            })
        else:
            return jsonify({'error': 'No question provided'}), 400
            
    except Exception as e:
        return jsonify({'error': f'API analysis failed: {str(e)}'}), 500

@app.route('/api-docs')
def api_documentation():
    """Serve API documentation page"""
    return render_template('api_documentation.html')

@app.route('/api/info')
def api_info():
    """Provide comprehensive API information in JSON format"""
    return jsonify({
        "name": "AutoStatIQ API",
        "version": "1.0.0",
        "description": "Statistical Analysis in Sec - Comprehensive statistical analysis API",
        "author": "Roman Chaudhary",
        "contact": "chaudharyroman.com.np",
        "base_url": request.host_url,
        "endpoints": {
            "health": "/health",
            "analysis": "/upload",
            "api_analysis": "/api/analyze",
            "download": "/download/<id>",
            "api_info": "/api/info",
            "documentation": "/api-docs"
        },
        "supported_formats": ["CSV", "XLSX", "DOCX", "PDF", "JSON", "TXT"],
        "file_size_limit": "16MB",
        "statistical_methods": [
            "Descriptive Statistics",
            "Frequency Analysis", 
            "Cross Tabulation & Chi-Square",
            "Correlation Analysis",
            "Regression Analysis",
            "Logistic Regression",
            "Hypothesis Testing",
            "ANOVA with Post-hoc Tests",
            "Control Charts (SPC)",
            "Normality Tests",
            "Outlier Detection",
            "Missing Data Analysis",
            "Principal Component Analysis",
            "Clustering Analysis",
            "Factor Analysis",
            "Survival Analysis",
            "Time Series Analysis",
            "Reliability Analysis"
        ],
        "control_chart_types": [
            "Individual (I) Charts",
            "Moving Range (MR) Charts", 
            "X-bar Charts",
            "R Charts",
            "C Charts",
            "P Charts",
            "U Charts"
        ],
        "features": [
            "file_upload",
            "statistical_analysis",
            "intelligent_interpretation",
            "report_generation",
            "control_charts",
            "advanced_statistics",
            "visualization_generation"
        ]
    })

@app.route('/health')
def health_check():
    """API health check endpoint"""
    try:
        # Check OpenAI configuration
        openai_configured = bool(openai.api_key)
        
        # Check if required modules are available
        required_modules = [
            'pandas', 'numpy', 'scipy', 'sklearn', 'matplotlib', 
            'seaborn', 'plotly', 'openai', 'docx'
        ]
        
        missing_modules = []
        for module in required_modules:
            try:
                __import__(module)
            except ImportError:
                missing_modules.append(module)
        
        status = "healthy" if not missing_modules else "unhealthy"
        
        return jsonify({
            "status": status,
            "version": "1.0.0",
            "timestamp": datetime.now().isoformat(),
            "openai_configured": openai_configured,
            "missing_modules": missing_modules,
            "features": [
                "file_upload",
                "statistical_analysis", 
                "intelligent_interpretation",
                "report_generation",
                "control_charts",
                "advanced_statistics",
                "visualization_generation"
            ],
            "supported_formats": ["CSV", "XLSX", "DOCX", "PDF", "JSON", "TXT"],
            "max_file_size": "16MB"
        })
        
    except Exception as e:
        return jsonify({
            "status": "error",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }), 500

if __name__ == '__main__':
    app.run(debug=True)
